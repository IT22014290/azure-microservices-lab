"""
Builds "Samishka H T - IT22014290.docx" — Azure Microservices Lab report.
Run: python3 build_docx.py
"""

import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = os.path.dirname(os.path.abspath(__file__))
SHOTS = os.path.join(BASE, "Screenshots")
OUT   = os.path.join(BASE, "Samishka H T - IT22014290.docx")

# ── helpers ────────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def para_fmt(p, size=11, bold=False, color=None, align=None, space_before=0, space_after=6):
    fmt = p.paragraph_format
    fmt.space_before = Pt(space_before)
    fmt.space_after  = Pt(space_after)
    if align:
        p.alignment = align
    for run in p.runs:
        run.font.size = Pt(size)
        run.font.bold = bold
        if color:
            run.font.color.rgb = RGBColor(*color)

def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.runs[0] if p.runs else p.add_run(text)
    if level == 1:
        run.font.color.rgb = RGBColor(0, 70, 127)
        run.font.size = Pt(16)
    elif level == 2:
        run.font.color.rgb = RGBColor(0, 112, 192)
        run.font.size = Pt(13)
    elif level == 3:
        run.font.color.rgb = RGBColor(0, 112, 192)
        run.font.size = Pt(11)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    return p

def add_body(doc, text, bold=False, size=11, color=None, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(space_after)
    return p

def add_code(doc, text):
    """Grey monospace code block."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.left_indent  = Cm(0.5)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(30, 30, 30)
    # grey background via paragraph shading
    pPr  = p._p.get_or_add_pPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  'F2F2F2')
    pPr.append(shd)
    return p

def add_bullet(doc, text, level=0, size=10.5):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.size = Pt(size)
    p.paragraph_format.left_indent   = Cm(0.5 + level * 0.5)
    p.paragraph_format.space_before  = Pt(0)
    p.paragraph_format.space_after   = Pt(3)
    return p

def add_screenshot(doc, filename, caption, width=6.0):
    path = os.path.join(SHOTS, filename)
    if not os.path.exists(path):
        add_body(doc, f"[Screenshot not found: {filename}]", color=(200,0,0))
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(path, width=Inches(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = cap.add_run(caption)
    run2.font.size  = Pt(9)
    run2.font.italic = True
    run2.font.color.rgb = RGBColor(89, 89, 89)
    cap.paragraph_format.space_after = Pt(10)

def add_task_header(doc, number, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(f"Task {number} – {title}")
    run.font.size  = Pt(14)
    run.font.bold  = True
    run.font.color.rgb = RGBColor(255, 255, 255)
    # blue background shading on the paragraph
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  '00467F')
    pPr.append(shd)

def add_divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '0070C0')
    pBdr.append(bottom)
    pPr.append(pBdr)

# ── build document ─────────────────────────────────────────────────────────────

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ══════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════════════════════════════════════════

doc.add_paragraph()  # spacer

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("SLIIT")
r.font.size = Pt(18)
r.font.bold = True
r.font.color.rgb = RGBColor(0, 70, 127)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run("Department of Computer Software Engineering\nFaculty of Computing")
r2.font.size = Pt(12)
r2.font.bold = True

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = p3.add_run("Module: Current Trends in Software Engineering (SE4010)")
r3.font.size = Pt(11)

p4 = doc.add_paragraph()
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
r4 = p4.add_run("2026 | Semester 1")
r4.font.size = Pt(11)

doc.add_paragraph()

p5 = doc.add_paragraph()
p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
r5 = p5.add_run("☁  Azure Microservices Deployment Lab")
r5.font.size  = Pt(22)
r5.font.bold  = True
r5.font.color.rgb = RGBColor(0, 112, 192)

p6 = doc.add_paragraph()
p6.alignment = WD_ALIGN_PARAGRAPH.CENTER
r6 = p6.add_run("Lab Report")
r6.font.size   = Pt(14)
r6.font.italic = True
r6.font.color.rgb = RGBColor(0, 112, 192)

doc.add_paragraph()

# Info table
info_table = doc.add_table(rows=3, cols=2)
info_table.style = 'Table Grid'
info_table.alignment = WD_ALIGN_PARAGRAPH.CENTER

labels = ["Name", "Student ID", "GitHub Repository"]
values = [
    "Samishka H T",
    "IT22014290",
    "https://github.com/IT22014290/azure-microservices-lab",
]
for i, (lbl, val) in enumerate(zip(labels, values)):
    c0 = info_table.cell(i, 0)
    c1 = info_table.cell(i, 1)
    set_cell_bg(c0, "00467F")
    c0.paragraphs[0].clear()
    r = c0.paragraphs[0].add_run(lbl)
    r.font.bold  = True
    r.font.color.rgb = RGBColor(255, 255, 255)
    r.font.size  = Pt(10)
    c1.paragraphs[0].clear()
    r2 = c1.paragraphs[0].add_run(val)
    r2.font.size = Pt(10)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 1 — LAB OVERVIEW
# ══════════════════════════════════════════════════════════════════════════════

add_heading(doc, "1. Lab Overview", level=1)
add_body(doc, (
    "This lab demonstrates building and deploying a containerized microservices application on Microsoft Azure "
    "using industry-standard tools and best practices. The application follows a three-tier microservices "
    "architecture deployed entirely on Azure cloud infrastructure."
))

add_heading(doc, "Learning Objectives", level=2)
for obj in [
    "Authenticate with Azure via the CLI",
    "Provision cloud infrastructure (Resource Group, Container Registry)",
    "Build and push Docker images to a private registry",
    "Deploy a containerized gateway service using Azure Container Apps",
    "Deploy an internal products microservice using Azure Container Apps",
    "Host a static frontend with Azure Static Web Apps",
    "Verify the deployment end-to-end and clean up resources",
]:
    add_bullet(doc, obj)

add_heading(doc, "Architecture Overview", level=2)
add_body(doc, "The application consists of three main tiers communicating over HTTPS:")

add_code(doc, """Browser
  │
  ▼  HTTPS
Azure Static Web App  (React frontend – global CDN)
  │
  ▼  HTTPS / REACT_APP_API_URL
Gateway Service  (Azure Container App · Node.js · port 3000 · External ingress)
  │
  ▼  internal HTTP / PRODUCTS_SERVICE_URL
Products Service  (Azure Container App · Node.js · port 3001 · Internal ingress)

Both Container Apps pull images from:
Azure Container Registry  (sliitmicroregistry.azurecr.io)""")

add_body(doc, "Data flow:", bold=True, space_after=2)
for step in [
    "The browser loads the React app from Azure Static Web Apps (global CDN, free SSL).",
    "The React app calls the Gateway via the public FQDN set in REACT_APP_API_URL.",
    "The Gateway proxies /api/products/* requests to the Products Service using its internal FQDN.",
    "Products Service is internal — it is NOT reachable from the public internet.",
]:
    add_bullet(doc, step)

add_heading(doc, "Technology Stack", level=2)
tbl = doc.add_table(rows=5, cols=3)
tbl.style = 'Table Grid'
headers = ["Component", "Azure Service", "Technology"]
rows_data = [
    ("Frontend", "Azure Static Web Apps", "React 18 / Node.js 18"),
    ("API Gateway", "Azure Container Apps (External)", "Node.js / Express"),
    ("Products Service", "Azure Container Apps (Internal)", "Node.js / Express"),
    ("Image Registry", "Azure Container Registry (Basic)", "Docker"),
]
hrow = tbl.rows[0]
for i, h in enumerate(headers):
    hrow.cells[i].paragraphs[0].clear()
    set_cell_bg(hrow.cells[i], "0070C0")
    r = hrow.cells[i].paragraphs[0].add_run(h)
    r.font.bold = True
    r.font.color.rgb = RGBColor(255, 255, 255)
    r.font.size = Pt(10)
for ri, row_data in enumerate(rows_data):
    row = tbl.rows[ri + 1]
    for ci, val in enumerate(row_data):
        row.cells[ci].paragraphs[0].clear()
        r = row.cells[ci].paragraphs[0].add_run(val)
        r.font.size = Pt(10)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 2 — PROJECT STRUCTURE
# ══════════════════════════════════════════════════════════════════════════════

add_heading(doc, "2. Project Structure", level=1)
add_code(doc, """azure-microservices-lab/
├── .github/
│   └── workflows/
│       └── azure-static-web-apps-lemon-bay-015f9eb00.yml   ← GitHub Actions CI/CD
├── Screenshots/
│   ├── Gateway check.png
│   ├── Gateway check 2.png
│   ├── Gateway Health Response.png
│   ├── Fontend 1.png
│   ├── Frontend 2.png
│   ├── Github action tab.png
│   ├── az resource list.png
│   └── container app logs.png
├── frontend/
│   ├── public/index.html
│   └── src/
│       ├── App.js          ← React UI (gateway health, services, products)
│       ├── App.css
│       └── index.js
├── gateway/
│   ├── server.js           ← Express API gateway + proxy middleware
│   ├── Dockerfile
│   ├── .dockerignore
│   └── package.json
├── products-service/
│   ├── server.js           ← Express products catalog API
│   ├── Dockerfile
│   ├── .dockerignore
│   └── package.json
└── LAB_DOCUMENTATION.md""")

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 3 — SOURCE CODE
# ══════════════════════════════════════════════════════════════════════════════

add_heading(doc, "3. Source Code", level=1)

add_heading(doc, "3.1 Gateway Service  (gateway/Dockerfile)", level=2)
add_body(doc, (
    "The Gateway is a Node.js/Express application that acts as the single public entry point. "
    "It proxies /api/products/* to the internal Products Service and exposes health, info, and services endpoints."
))
add_code(doc, """FROM node:18-alpine
WORKDIR /app
COPY package*.json ./
RUN npm install --production
COPY . .
EXPOSE 3000
CMD ["node", "server.js"]""")

add_heading(doc, "gateway/server.js — key endpoints", level=3)
tbl2 = doc.add_table(rows=7, cols=3)
tbl2.style = 'Table Grid'
h2row = tbl2.rows[0]
for i, h in enumerate(["Method", "Path", "Description"]):
    set_cell_bg(h2row.cells[i], "0070C0")
    h2row.cells[i].paragraphs[0].clear()
    r = h2row.cells[i].paragraphs[0].add_run(h)
    r.font.bold = True
    r.font.color.rgb = RGBColor(255, 255, 255)
    r.font.size = Pt(10)
ep_rows = [
    ("GET", "/",               "Service info & endpoint list"),
    ("GET", "/health",         "Health check – returns status, uptime, timestamp"),
    ("GET", "/api/info",       "Gateway runtime info (region, Node version)"),
    ("GET", "/api/services",   "List all deployed services and their status"),
    ("GET", "/api/products",   "All products – proxied to Products Service"),
    ("GET", "/api/products/:id","Single product – proxied to Products Service"),
]
for ri, (m, path, desc) in enumerate(ep_rows):
    row = tbl2.rows[ri + 1]
    for ci, val in enumerate([m, path, desc]):
        row.cells[ci].paragraphs[0].clear()
        r = row.cells[ci].paragraphs[0].add_run(val)
        r.font.size = Pt(9.5)
        if ci == 0:
            r.font.color.rgb = RGBColor(0, 112, 192)
            r.font.bold = True

doc.add_paragraph()

add_heading(doc, "3.2 Products Service  (products-service/Dockerfile)", level=2)
add_body(doc, (
    "A standalone Node.js/Express microservice that serves a mock Azure-services product catalog. "
    "Deployed with internal ingress — only reachable from within the Container Apps environment."
))
add_code(doc, """FROM node:18-alpine
WORKDIR /app
COPY package*.json ./
RUN npm install --production
COPY . .
EXPOSE 3001
CMD ["node", "server.js"]""")

add_heading(doc, "Products Service endpoints", level=3)
tbl3 = doc.add_table(rows=4, cols=3)
tbl3.style = 'Table Grid'
h3row = tbl3.rows[0]
for i, h in enumerate(["Method", "Path", "Description"]):
    set_cell_bg(h3row.cells[i], "0070C0")
    h3row.cells[i].paragraphs[0].clear()
    r = h3row.cells[i].paragraphs[0].add_run(h)
    r.font.bold = True
    r.font.color.rgb = RGBColor(255, 255, 255)
    r.font.size = Pt(10)
ps_rows = [
    ("GET", "/health",        "Health check"),
    ("GET", "/products",      "List all products (supports ?category= filter)"),
    ("GET", "/products/:id",  "Get single product by ID"),
]
for ri, (m, path, desc) in enumerate(ps_rows):
    row = tbl3.rows[ri + 1]
    for ci, val in enumerate([m, path, desc]):
        row.cells[ci].paragraphs[0].clear()
        r = row.cells[ci].paragraphs[0].add_run(val)
        r.font.size = Pt(9.5)
        if ci == 0:
            r.font.color.rgb = RGBColor(0, 112, 192)
            r.font.bold = True

doc.add_paragraph()

add_heading(doc, "3.3 Frontend  (React 18)", level=2)
add_body(doc, (
    "React 18 single-page application hosted on Azure Static Web Apps. Reads REACT_APP_API_URL at "
    "build time to know the gateway address. Displays: gateway health status, deployed services list, "
    "product catalog, and gateway runtime info."
))

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 4 — CI/CD PIPELINE
# ══════════════════════════════════════════════════════════════════════════════

add_heading(doc, "4. CI/CD Pipeline — GitHub Actions", level=1)
add_body(doc, (
    "Azure Static Web Apps automatically creates a GitHub Actions workflow that builds and deploys "
    "the React frontend on every push to main."
))
add_code(doc, """# .github/workflows/azure-static-web-apps-lemon-bay-015f9eb00.yml
name: Azure Static Web Apps CI/CD

on:
  push:
    branches: [ main ]
  pull_request:
    types: [opened, synchronize, reopened, closed]
    branches: [ main ]

jobs:
  build_and_deploy_job:
    if: github.event_name == 'push' || (github.event_name == 'pull_request' && github.event.action != 'closed')
    runs-on: ubuntu-latest
    steps:
      - uses: actions/checkout@v3
        with:
          submodules: true

      - name: Build And Deploy
        uses: Azure/static-web-apps-deploy@v1
        with:
          azure_static_web_apps_api_token: ${{ secrets.AZURE_STATIC_WEB_APPS_API_TOKEN_LEMON_BAY_015F9EB00 }}
          repo_token: ${{ secrets.GITHUB_TOKEN }}
          action: "upload"
          app_location: "/frontend"
          output_location: "build"
        env:
          REACT_APP_API_URL: https://gateway.redwave-597fbb8a.eastasia.azurecontainerapps.io

  close_pull_request_job:
    if: github.event_name == 'pull_request' && github.event.action == 'closed'
    runs-on: ubuntu-latest
    steps:
      - name: Close Pull Request
        uses: Azure/static-web-apps-deploy@v1
        with:
          azure_static_web_apps_api_token: ${{ secrets.AZURE_STATIC_WEB_APPS_API_TOKEN_LEMON_BAY_015F9EB00 }}
          action: "close" """)

add_body(doc, "How the pipeline works:", bold=True, space_after=2)
for step in [
    "On every push to main, GitHub Actions checks out the repo.",
    "The React app in /frontend is built with REACT_APP_API_URL injected at build time.",
    "The build/ output is deployed to Azure Static Web Apps CDN globally.",
    "On PR close, the preview deployment is automatically cleaned up.",
]:
    add_bullet(doc, step)

add_screenshot(doc, "Github action tab.png",
    "Figure 1 — GitHub Actions CI/CD pipeline running successfully for the Azure Static Web Apps deployment", 6.0)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 5 — AZURE DEPLOYMENT (Tasks 1-6)
# ══════════════════════════════════════════════════════════════════════════════

add_heading(doc, "5. Azure Deployment — Step by Step", level=1)

add_body(doc, "Environment variables set once before running deployment commands:", bold=True)
add_code(doc, """REGISTRY=sliitmicroregistry   # ACR name (globally unique)
RG=microservices-rg
LOCATION=eastus""")

add_divider(doc)

# ─── TASK 1 ───────────────────────────────────────────────────────────────────
add_task_header(doc, 1, "Login to Azure")

add_heading(doc, "Objective", level=3)
add_body(doc, (
    "Authenticate the terminal session with Microsoft Azure and confirm the correct subscription is active. "
    "All subsequent tasks depend on this session."
))

add_heading(doc, "Background", level=3)
add_body(doc, (
    "The Azure CLI (az) uses OAuth2 browser-based authentication to obtain an access token. "
    "Once authenticated, the token is cached locally and reused for subsequent commands."
))

add_heading(doc, "Commands", level=3)

add_body(doc, "Step 1.1 — Verify Azure CLI is installed:", bold=True, space_after=2)
add_code(doc, "az --version")

add_body(doc, "Step 1.2 — Authenticate with Azure (browser-based login):", bold=True, space_after=2)
add_code(doc, "az login")

add_body(doc, "Step 1.3 — Verify active account and subscription:", bold=True, space_after=2)
add_code(doc, "az account show")

add_body(doc, "Step 1.4 — (Optional) Set a specific subscription:", bold=True, space_after=2)
add_code(doc, """az account list --output table
az account set --subscription "<Subscription Name or ID>" """)

add_heading(doc, "Verification Checklist", level=3)
for item in [
    "az --version shows version 2.50 or higher",
    "az login completed without errors",
    "az account show shows correct email address",
    "Subscription state is Enabled",
]:
    add_bullet(doc, item)

add_divider(doc)

# ─── TASK 2 ───────────────────────────────────────────────────────────────────
add_task_header(doc, 2, "Create Resource Group & Container Registry")

add_heading(doc, "Objective", level=3)
add_body(doc, (
    "Create an Azure Resource Group to logically contain all lab resources and provision an "
    "Azure Container Registry (ACR) to store Docker images privately."
))

add_heading(doc, "Background", level=3)
add_body(doc, (
    "A Resource Group is a logical container for related Azure resources — they share the same lifecycle "
    "and can be managed and deleted together. Azure Container Registry (ACR) is a private Docker registry "
    "that integrates natively with Azure Container Apps."
))

add_heading(doc, "Commands", level=3)

add_body(doc, "Step 2.1 — Create the Resource Group:", bold=True, space_after=2)
add_code(doc, "az group create --name microservices-rg --location eastus")

add_body(doc, "Expected output:", space_after=2)
add_code(doc, """{
  "id": "/subscriptions/.../resourceGroups/microservices-rg",
  "location": "eastus",
  "name": "microservices-rg",
  "properties": { "provisioningState": "Succeeded" }
}""")

add_body(doc, "Step 2.2 — Create the Azure Container Registry (Basic SKU):", bold=True, space_after=2)
add_code(doc, """az acr create \\
  --resource-group microservices-rg \\
  --name sliitmicroregistry \\
  --sku Basic""")

add_body(doc, "Step 2.3 — Authenticate Docker with ACR:", bold=True, space_after=2)
add_code(doc, "az acr login --name sliitmicroregistry")
add_body(doc, "Expected: Login Succeeded", color=(0, 128, 0))

add_heading(doc, "Verification Checklist", level=3)
for item in [
    "Resource group microservices-rg shows provisioningState: Succeeded",
    "Container registry sliitmicroregistry created successfully",
    "az acr login returns Login Succeeded",
    "Registry visible in Azure Portal under microservices-rg",
]:
    add_bullet(doc, item)

add_divider(doc)

# ─── TASK 3 ───────────────────────────────────────────────────────────────────
add_task_header(doc, 3, "Build & Push Docker Images")

add_heading(doc, "Objective", level=3)
add_body(doc, (
    "Build Docker images for the gateway and products microservices and push them to the Azure "
    "Container Registry. These images are used by Container Apps to run the services in the cloud."
))

add_heading(doc, "Background", level=3)
add_body(doc, (
    "Docker images are portable, self-contained packages that include application code and all runtime "
    "dependencies. By tagging images with the ACR login server address (sliitmicroregistry.azurecr.io), "
    "Docker knows where to push them."
))

add_heading(doc, "Commands", level=3)

add_body(doc, "Step 3.1 — Review gateway Dockerfile:", bold=True, space_after=2)
add_code(doc, """# gateway/Dockerfile
FROM node:18-alpine
WORKDIR /app
COPY package*.json ./
RUN npm install --production
COPY . .
EXPOSE 3000
CMD ["node", "server.js"]""")

add_body(doc, "Step 3.2 — Build the gateway image:", bold=True, space_after=2)
add_code(doc, "docker build -t sliitmicroregistry.azurecr.io/gateway:v1 ./gateway")

add_body(doc, "Step 3.3 — Push the gateway image to ACR:", bold=True, space_after=2)
add_code(doc, "docker push sliitmicroregistry.azurecr.io/gateway:v1")

add_body(doc, "Step 3.4 — Build and push the products-service image:", bold=True, space_after=2)
add_code(doc, """docker build -t sliitmicroregistry.azurecr.io/products-service:v1 ./products-service
docker push  sliitmicroregistry.azurecr.io/products-service:v1""")

add_body(doc, "Step 3.5 — Verify images in ACR:", bold=True, space_after=2)
add_code(doc, """az acr repository list --name sliitmicroregistry --output table
az acr repository show-tags --name sliitmicroregistry --repository gateway --output table""")

add_body(doc, "Expected output:", space_after=2)
add_code(doc, """Result
-----------------
gateway
products-service""")

add_heading(doc, "Docker Desktop — Image Layers", level=3)
add_screenshot(doc, "Screenshot 2026-03-14 at 13.57.18.png",
    "Figure 2 — Docker Desktop showing the gateway:v1 image pushed to sliitmicroregistry.azurecr.io with 15 layers", 6.0)

add_heading(doc, "Docker Desktop — Running Containers", level=3)
add_screenshot(doc, "Screenshot 2026-03-14 at 13.56.34.png",
    "Figure 3 — Docker Desktop Containers view showing gateway:v1 container running locally during testing", 6.0)

add_screenshot(doc, "Screenshot 2026-03-14 at 13.56.54.png",
    "Figure 4 — Docker Desktop Logs view showing 'Gateway service running on port 3000' confirmation", 6.0)

add_heading(doc, "Verification Checklist", level=3)
for item in [
    "docker build completed without errors for both services",
    "docker push completed successfully — all layers uploaded",
    "az acr repository list shows both gateway and products-service",
    "Tag v1 is visible in both repositories",
]:
    add_bullet(doc, item)

add_divider(doc)

# ─── TASK 4 ───────────────────────────────────────────────────────────────────
add_task_header(doc, 4, "Deploy Container Apps")

add_heading(doc, "Objective", level=3)
add_body(doc, (
    "Register the Azure Container Apps provider, create a managed environment, and deploy both the "
    "gateway (external) and products-service (internal) as Azure Container Apps."
))

add_heading(doc, "Background", level=3)
add_body(doc, (
    "Azure Container Apps is a serverless platform for running containerized applications. Unlike AKS, "
    "it abstracts away infrastructure management. Container Apps run within an Environment, which provides "
    "shared networking and observability. Services with internal ingress are only reachable from within "
    "the same environment."
))

add_heading(doc, "Commands", level=3)

add_body(doc, "Step 4.1 — Register required Azure resource providers:", bold=True, space_after=2)
add_code(doc, """az provider register --namespace Microsoft.App --wait
az provider register --namespace Microsoft.OperationalInsights --wait""")

add_body(doc, "Step 4.2 — Create the Container Apps managed environment:", bold=True, space_after=2)
add_code(doc, """az containerapp env create \\
  --name micro-env \\
  --resource-group microservices-rg \\
  --location eastus""")

add_body(doc, "Step 4.3 — Enable ACR admin credentials:", bold=True, space_after=2)
add_code(doc, """az acr update -n sliitmicroregistry --admin-enabled true

# Retrieve password
ACR_PASSWORD=$(az acr credential show --name sliitmicroregistry --query "passwords[0].value" -o tsv)""")

add_body(doc, "Step 4.4 — Deploy products-service with INTERNAL ingress:", bold=True, space_after=2)
add_code(doc, """az containerapp create \\
  --name products-service \\
  --resource-group microservices-rg \\
  --environment micro-env \\
  --image sliitmicroregistry.azurecr.io/products-service:v1 \\
  --target-port 3001 \\
  --ingress internal \\
  --registry-server sliitmicroregistry.azurecr.io \\
  --registry-username sliitmicroregistry \\
  --registry-password $ACR_PASSWORD""")

add_body(doc, "Step 4.5 — Retrieve the internal FQDN of products-service:", bold=True, space_after=2)
add_code(doc, """PRODUCTS_FQDN=$(az containerapp show \\
  --name products-service \\
  --resource-group microservices-rg \\
  --query properties.configuration.ingress.fqdn \\
  --output tsv)

echo "Products internal URL: https://$PRODUCTS_FQDN" """)

add_body(doc, "Step 4.6 — Deploy gateway with EXTERNAL ingress:", bold=True, space_after=2)
add_code(doc, """az containerapp create \\
  --name gateway \\
  --resource-group microservices-rg \\
  --environment micro-env \\
  --image sliitmicroregistry.azurecr.io/gateway:v1 \\
  --target-port 3000 \\
  --ingress external \\
  --registry-server sliitmicroregistry.azurecr.io \\
  --registry-username sliitmicroregistry \\
  --registry-password $ACR_PASSWORD \\
  --env-vars PRODUCTS_SERVICE_URL=https://$PRODUCTS_FQDN""")

add_body(doc, "Step 4.7 — Retrieve the public gateway URL:", bold=True, space_after=2)
add_code(doc, """GATEWAY_FQDN=$(az containerapp show \\
  --name gateway \\
  --resource-group microservices-rg \\
  --query properties.configuration.ingress.fqdn \\
  --output tsv)

echo "Gateway URL: https://$GATEWAY_FQDN" """)

add_body(doc, "Key parameters explained:", bold=True, space_after=2)
for item in [
    "--target-port 3000 / 3001 — the port the application listens on inside the container",
    "--ingress external — makes the app publicly accessible via HTTPS (gateway only)",
    "--ingress internal — service only reachable from within micro-env (products-service)",
    "--env-vars PRODUCTS_SERVICE_URL — passes the products-service internal URL to gateway",
]:
    add_bullet(doc, item)

add_heading(doc, "Verification Checklist", level=3)
for item in [
    "Microsoft.App provider registered successfully",
    "Container Apps Environment micro-env created",
    "products-service deployed with internal ingress",
    "gateway deployed with external ingress",
    "Public FQDN URL returned by az containerapp show",
    "Gateway service accessible in browser via HTTPS",
]:
    add_bullet(doc, item)

add_divider(doc)

# ─── TASK 5 ───────────────────────────────────────────────────────────────────
add_task_header(doc, 5, "Deploy Static Web App Frontend")

add_heading(doc, "Objective", level=3)
add_body(doc, (
    "Deploy the React frontend using Azure Static Web Apps — a fully managed hosting service with "
    "global CDN distribution, free SSL/TLS certificates, and automatic CI/CD via GitHub Actions."
))

add_heading(doc, "Commands", level=3)

add_body(doc, "Step 5.1 — Create the Static Web App resource:", bold=True, space_after=2)
add_code(doc, """az staticwebapp create \\
  --name sliit-frontend-app \\
  --resource-group microservices-rg \\
  --location eastus \\
  --source https://github.com/IT22014290/azure-microservices-lab \\
  --branch main \\
  --app-location "/frontend" \\
  --output-location "build" """)

add_body(doc, "Step 5.2 — Point the frontend at the gateway URL:", bold=True, space_after=2)
add_code(doc, """az staticwebapp appsettings set \\
  --name sliit-frontend-app \\
  --resource-group microservices-rg \\
  --setting-names REACT_APP_API_URL=https://gateway.redwave-597fbb8a.eastasia.azurecontainerapps.io""")

add_body(doc, "Step 5.3 — Get the public frontend URL:", bold=True, space_after=2)
add_code(doc, """az staticwebapp show \\
  --name sliit-frontend-app \\
  --resource-group microservices-rg \\
  --query defaultHostname \\
  --output tsv""")

add_body(doc, "Step 5.4 — Monitor the GitHub Actions deployment:", bold=True, space_after=2)
for step in [
    "Go to the GitHub repository in the browser",
    "Click the Actions tab",
    "Watch the Azure Static Web Apps CI/CD workflow run",
    "Wait for the green checkmark (usually 2–5 minutes)",
]:
    add_bullet(doc, step)

add_heading(doc, "Verification Checklist", level=3)
for item in [
    "Static Web App resource created successfully in Azure",
    "GitHub Actions workflow ran and completed with green checkmark",
    "Frontend accessible via assigned Azure URL",
    "Page renders without errors in browser console",
]:
    add_bullet(doc, item)

add_divider(doc)

# ─── TASK 6 ───────────────────────────────────────────────────────────────────
add_task_header(doc, 6, "Verify Deployment & Cleanup")

add_heading(doc, "Objective", level=3)
add_body(doc, (
    "Perform end-to-end verification of the entire deployment to confirm all services are running, "
    "then delete all Azure resources to avoid incurring ongoing costs."
))

add_heading(doc, "Verification Commands", level=3)

add_body(doc, "Step 6.1 — List all resources in the resource group:", bold=True, space_after=2)
add_code(doc, "az resource list --resource-group microservices-rg --output table")

add_body(doc, "Expected resources:", space_after=2)
for item in [
    "Container Registry — sliitmicroregistry (Microsoft.ContainerRegistry/registries)",
    "Container Apps Environment — micro-env (Microsoft.App/managedEnvironments)",
    "Container App — gateway (Microsoft.App/containerApps)",
    "Container App — products-service (Microsoft.App/containerApps)",
    "Static Web App — sliit-frontend-app (Microsoft.Web/staticSites)",
]:
    add_bullet(doc, item)

add_body(doc, "Step 6.2 — Test the gateway health endpoint:", bold=True, space_after=2)
add_code(doc, """curl https://gateway.redwave-597fbb8a.eastasia.azurecontainerapps.io/health
# Expected: { "status": "ok", "uptime": ..., "productsServiceConfigured": true }""")

add_body(doc, "Step 6.3 — Test the products endpoint (proxied through gateway):", bold=True, space_after=2)
add_code(doc, "curl https://gateway.redwave-597fbb8a.eastasia.azurecontainerapps.io/api/products")

add_body(doc, "Step 6.4 — Review Container App logs:", bold=True, space_after=2)
add_code(doc, """az containerapp logs show \\
  --name gateway \\
  --resource-group microservices-rg \\
  --follow false

az containerapp logs show \\
  --name products-service \\
  --resource-group microservices-rg \\
  --follow false""")

add_body(doc, "Step 6.5 — Cleanup — delete ALL resources:", bold=True, space_after=2)
add_code(doc, "az group delete --name microservices-rg --yes")
add_body(doc, (
    "Warning: The --yes flag skips the confirmation prompt. This operation is irreversible and deletes "
    "all resources including the Container Registry, Container Apps, and Static Web App."
), color=(192, 0, 0))

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 6 — SCREENSHOTS
# ══════════════════════════════════════════════════════════════════════════════

doc.add_page_break()
add_heading(doc, "6. Deployment Screenshots", level=1)
add_body(doc, (
    "The following screenshots document the live deployment and verification steps completed during the lab."
))

add_heading(doc, "6.1 Gateway Root Endpoint", level=2)
add_screenshot(doc, "Gateway check.png",
    "Figure 5 — Gateway root endpoint (/) returning service info, version 2.0.0, and available endpoints list", 6.0)

add_heading(doc, "6.2 Gateway Endpoint — Additional Verification", level=2)
add_screenshot(doc, "Gateway check 2.png",
    "Figure 6 — Additional gateway endpoint check showing services and API info responses", 6.0)

add_heading(doc, "6.3 Gateway Health Endpoint", level=2)
add_screenshot(doc, "Gateway Health Response.png",
    'Figure 7 — Gateway /health endpoint returning { "status": "ok", "productsServiceConfigured": true } via HTTPS', 6.0)

add_heading(doc, "6.4 Azure Resource List", level=2)
add_screenshot(doc, "az resource list.png",
    "Figure 8 — Terminal output of 'az resource list --resource-group microservices-rg --output table' showing all provisioned Azure resources", 6.0)

add_heading(doc, "6.5 Container App Logs", level=2)
add_screenshot(doc, "container app logs.png",
    "Figure 9 — Gateway Container App logs showing incoming HTTP requests being proxied to the products-service", 6.0)

add_heading(doc, "6.6 React Frontend — View 1", level=2)
add_screenshot(doc, "Fontend 1.png",
    "Figure 10 — React frontend loaded from Azure Static Web Apps showing gateway health status and deployed services", 6.0)

add_heading(doc, "6.7 React Frontend — View 2", level=2)
add_screenshot(doc, "Frontend 2.png",
    "Figure 11 — React frontend displaying the product catalog (6 Azure services) fetched through the gateway", 6.0)

add_heading(doc, "6.8 GitHub Actions CI/CD Workflow", level=2)
add_screenshot(doc, "Github action tab.png",
    "Figure 12 — GitHub Actions tab showing the Azure Static Web Apps CI/CD pipeline completing successfully", 6.0)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 7 — SUMMARY
# ══════════════════════════════════════════════════════════════════════════════

doc.add_page_break()
add_heading(doc, "7. Summary", level=1)

tbl_s = doc.add_table(rows=5, cols=4)
tbl_s.style = 'Table Grid'
hs = tbl_s.rows[0]
for i, h in enumerate(["Component", "Azure Service", "Ingress", "Port"]):
    set_cell_bg(hs.cells[i], "00467F")
    hs.cells[i].paragraphs[0].clear()
    r = hs.cells[i].paragraphs[0].add_run(h)
    r.font.bold = True
    r.font.color.rgb = RGBColor(255, 255, 255)
    r.font.size = Pt(10)
sum_rows = [
    ("Frontend",         "Azure Static Web Apps",  "Public (CDN)",      "443"),
    ("Gateway",          "Azure Container Apps",    "External (public)", "3000 → 443"),
    ("Products Service", "Azure Container Apps",    "Internal only",     "3001"),
    ("Image Registry",   "Azure Container Registry","N/A",               "N/A"),
]
for ri, row_data in enumerate(sum_rows):
    row = tbl_s.rows[ri + 1]
    for ci, val in enumerate(row_data):
        row.cells[ci].paragraphs[0].clear()
        r = row.cells[ci].paragraphs[0].add_run(val)
        r.font.size = Pt(10)

doc.add_paragraph()
add_body(doc, "Key Takeaways:", bold=True, space_after=2)
for kt in [
    "The API Gateway pattern centralizes routing, hiding internal services from the public internet.",
    "Internal ingress on the Products Service ensures it is only reachable by the gateway within micro-env.",
    "GitHub Actions automates frontend deployment — any push to main triggers a rebuild and re-deploy.",
    "Azure Container Registry provides secure, private image storage with native Container Apps integration.",
    "The entire infrastructure can be torn down with a single az group delete command.",
]:
    add_bullet(doc, kt)

add_heading(doc, "Useful Commands Reference", level=2)
add_code(doc, """# Check login status
az account show

# List all resources in group
az resource list --resource-group microservices-rg --output table

# View container app details
az containerapp show --name gateway --resource-group microservices-rg

# Stream live logs
az containerapp logs show --name gateway --resource-group microservices-rg --follow true

# Restart container app
az containerapp revision restart --name gateway --resource-group microservices-rg \\
  --revision <revision-name>

# List ACR images
az acr repository list --name sliitmicroregistry --output table

# Test gateway health
curl https://gateway.redwave-597fbb8a.eastasia.azurecontainerapps.io/health

# Delete all resources (cleanup)
az group delete --name microservices-rg --yes""")

# ══════════════════════════════════════════════════════════════════════════════
# TROUBLESHOOTING
# ══════════════════════════════════════════════════════════════════════════════

add_heading(doc, "8. Troubleshooting Guide", level=1)

tbl_t = doc.add_table(rows=9, cols=3)
tbl_t.style = 'Table Grid'
ht = tbl_t.rows[0]
for i, h in enumerate(["Problem", "Possible Cause", "Solution"]):
    set_cell_bg(ht.cells[i], "00467F")
    ht.cells[i].paragraphs[0].clear()
    r = ht.cells[i].paragraphs[0].add_run(h)
    r.font.bold = True
    r.font.color.rgb = RGBColor(255, 255, 255)
    r.font.size = Pt(10)
trouble_rows = [
    ("az login not working",        "Session expired",                           "Run az login again; headless: az login --use-device-code"),
    ("docker push fails",           "ACR login missing or expired",              "Run az acr login --name sliitmicroregistry"),
    ("Container not accessible",    "Port not exposed or ingress not set",       "Verify --target-port matches EXPOSE and --ingress external"),
    ("Static web app deploy fails", "GitHub Actions auth issue",                 "Re-authenticate: az staticwebapp secrets reset-api-key"),
    ("Container app returns 503",   "Image pull failure or app startup error",   "Check logs: az containerapp logs show; verify ACR credentials"),
    ("Provider registration error", "Insufficient subscription permissions",     "Contact subscription admin to register Microsoft.App namespace"),
    ("Resource already exists",     "Duplicate name in same region",             "Append unique suffix (e.g. student ID) to resource name"),
    ("ACR push: denied",            "Admin credentials not enabled",             "Run az acr update -n sliitmicroregistry --admin-enabled true"),
]
for ri, row_data in enumerate(trouble_rows):
    row = tbl_t.rows[ri + 1]
    for ci, val in enumerate(row_data):
        row.cells[ci].paragraphs[0].clear()
        r = row.cells[ci].paragraphs[0].add_run(val)
        r.font.size = Pt(9)
        if ci == 0:
            r.font.color.rgb = RGBColor(192, 0, 0)
            r.font.bold = True

# ══════════════════════════════════════════════════════════════════════════════
# SAVE
# ══════════════════════════════════════════════════════════════════════════════

doc.save(OUT)
print(f"Saved → {OUT}")
